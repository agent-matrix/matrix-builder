"""Batch 3 — POST /api/v1/ingest/document → deterministic ProjectBrief."""
from __future__ import annotations

import io

from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Client Delivery Portal", level=1)
    doc.add_paragraph("A secure portal for consulting clients with files and status tracking.")
    doc.add_heading("Features", level=2)
    doc.add_paragraph("File uploads", style="List Bullet")
    doc.add_paragraph("Status tracking", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_ingest_docx_builds_brief():
    resp = client.post(
        "/api/v1/ingest/document",
        files={"file": ("brief.docx", _docx_bytes(), _DOCX_MIME)},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["source_type"] == "document"
    brief = data["brief"]
    assert brief["title"] == "Client Delivery Portal"
    assert "secure portal" in brief["summary"].lower()
    assert any("upload" in f.lower() for f in brief["features"])
    assert brief["enhanced_by"] == "deterministic"
    assert "Client Delivery Portal" in data["idea"]


def test_ingest_markdown_builds_brief():
    md = b"# Inventory Forecasting\n\nForecast stock levels for retail.\n\n## Features\n- demand forecast\n- reorder alerts\n"
    resp = client.post("/api/v1/ingest/document", files={"file": ("brief.md", md, "text/markdown")})
    assert resp.status_code == 200, resp.text
    assert resp.json()["brief"]["title"] == "Inventory Forecasting"
    assert any("forecast" in f.lower() for f in resp.json()["brief"]["features"])


def test_ingest_unsupported_format_rejected():
    resp = client.post("/api/v1/ingest/document", files={"file": ("design.png", b"\x89PNG\r\n", "image/png")})
    assert resp.status_code == 415


def test_ingest_empty_rejected():
    resp = client.post("/api/v1/ingest/document", files={"file": ("brief.md", b"", "text/markdown")})
    assert resp.status_code == 400


def _valid_blueprint() -> dict:
    return {
        "blueprint_id": "bp-user-001", "candidate_id": "user", "name": "My Custom App",
        "slug": "my-custom-app", "idea": "user-provided spec", "quality_level": "standard",
        "stack": {"frontend": "nextjs", "backend": "fastapi"},
        "pages": ["/"], "services": ["api"], "api_routes": [],
        "required_files": ["README.md", "MATRIX_BLUEPRINT.yaml"],
        "allowed_change_roots": ["backend/"], "forbidden_changes": [],
        "tasks": [{"task_id": "TASK-001", "title": "Implement it",
                   "allowed_files": ["backend/app.py"], "acceptance_criteria": ["works"]}],
        "acceptance_commands": ["pytest -q"], "standards_lock_ref": "MATRIX_STANDARDS.lock",
    }


def test_blueprint_import_valid_forces_control_files_forbidden():
    resp = client.post("/api/v1/ingest/blueprint", json={"blueprint": _valid_blueprint()})
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["valid"] is True and not data["errors"]
    assert "MATRIX_BLUEPRINT.yaml" in data["blueprint"]["forbidden_changes"]
    assert "MATRIX_STANDARDS.lock" in data["blueprint"]["forbidden_changes"]


def test_blueprint_import_invalid_schema_reports_errors():
    resp = client.post("/api/v1/ingest/blueprint", json={"blueprint": {"name": "x"}})
    assert resp.status_code == 200
    assert resp.json()["valid"] is False and resp.json()["errors"]


def test_blueprint_import_requires_tasks():
    bp = _valid_blueprint(); bp["tasks"] = []
    resp = client.post("/api/v1/ingest/blueprint", json={"blueprint": bp})
    assert resp.json()["valid"] is False


def test_bundle_from_blueprint_skips_ai():
    resp = client.post("/api/v1/bundles", json={"blueprint": _valid_blueprint(), "preferred_coder": "claude-code"})
    assert resp.status_code == 200, resp.text
    assert resp.json()["bundle_id"]
