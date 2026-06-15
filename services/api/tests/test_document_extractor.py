"""Unit tests for the lightweight PDF/DOCX → markdown extractor (Path B)."""
from __future__ import annotations

import io

import pytest

from app.services.document_extractor import (
    DocumentExtractionError,
    UnsupportedDocumentError,
    extract_markdown,
)


def _make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Consulting Delivery Portal", level=1)
    doc.add_paragraph("A client portal for delivery teams with audit history.")
    doc.add_heading("Features", level=2)
    doc.add_paragraph("Document uploads", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Access"
    table.rows[1].cells[0].text = "Admin"
    table.rows[1].cells[1].text = "Full"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf(text: str = "Hello Matrix Brief") -> bytes:
    """Assemble a minimal but valid one-page PDF (correct xref offsets) with extractable text."""
    stream = b"BT /F1 24 Tf 72 700 Td (" + text.encode() + b") Tj ET"
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>",
        b"<</Length " + str(len(stream)).encode() + b">>stream\n" + stream + b"\nendstream",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for i, body in enumerate(objs, 1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj" + body + b"endobj\n"
    xref_pos = len(out)
    n = len(objs) + 1
    out += b"xref\n0 " + str(n).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        out += ("%010d 00000 n \n" % off).encode()
    out += b"trailer<</Size " + str(n).encode() + b"/Root 1 0 R>>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    return bytes(out)


def test_docx_extracts_headings_paragraphs_and_table():
    md = extract_markdown("brief.docx", _make_docx())
    assert "# Consulting Delivery Portal" in md
    assert "## Features" in md
    assert "client portal for delivery teams" in md
    assert "| Role | Access |" in md
    assert "| Admin | Full |" in md


def test_pdf_extracts_text():
    md = extract_markdown("brief.pdf", _make_pdf())
    assert "Hello Matrix Brief" in md


def test_markdown_and_txt_pass_through():
    assert "# Goals" in extract_markdown("brief.md", b"# Goals\n- ship it")
    assert "plain notes" in extract_markdown("notes.txt", b"plain notes")


def test_unsupported_format_rejected():
    # Images (design path) and JSON (skip-AI import path) are handled elsewhere, not here.
    for name in ("design.png", "photo.jpg", "blueprint.json", "noext"):
        with pytest.raises(UnsupportedDocumentError):
            extract_markdown(name, b"whatever")


def test_corrupt_pdf_raises_extraction_error():
    with pytest.raises(DocumentExtractionError):
        extract_markdown("broken.pdf", b"not a real pdf at all")


def test_output_is_capped():
    from app.services.document_extractor import MAX_MARKDOWN_CHARS

    big = io.BytesIO()
    from docx import Document

    doc = Document()
    for _ in range(20000):
        doc.add_paragraph("x" * 40)
    doc.save(big)
    md = extract_markdown("big.docx", big.getvalue())
    assert len(md) <= MAX_MARKDOWN_CHARS
