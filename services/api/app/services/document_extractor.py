"""Lightweight document → Markdown extraction for the brief upload path (Path B).

PDF and DOCX only, on purpose. Uses two tiny pure-Python libraries (no torch/ML, no OCR), so it runs
in-process on the cpu-basic Space — Docling was evaluated and rejected as too heavy. Scanned/image
PDFs yield little text (no OCR); that is acceptable for v1. The extracted markdown becomes the input
to a ProjectBrief; it never bypasses the deterministic engine or the Matrix contract.
"""
from __future__ import annotations

import io
from pathlib import PurePosixPath

# Text-bearing document formats this extractor understands. Images (PNG/JPG) are a separate
# design path (OllaBridge vision); a complete Blueprint JSON is the skip-AI import path — neither
# goes through here. Detection is by extension; the API layer should also check MIME.
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".md", ".markdown", ".txt")

# Guardrail: cap extracted text so a huge upload can't blow up downstream payloads.
MAX_MARKDOWN_CHARS = 200_000


class UnsupportedDocumentError(ValueError):
    """Raised when an uploaded file is not a PDF or DOCX."""


class DocumentExtractionError(RuntimeError):
    """Raised when a supported file cannot be parsed (corrupt, encrypted, or empty)."""


def extract_markdown(filename: str, data: bytes) -> str:
    """Convert an uploaded PDF or DOCX (raw bytes) into plain Markdown text.

    Raises UnsupportedDocumentError for any other type, DocumentExtractionError on parse failure.
    """
    ext = PurePosixPath(filename or "").suffix.lower()
    if ext == ".pdf":
        text = _pdf_to_markdown(data)
    elif ext == ".docx":
        text = _docx_to_markdown(data)
    elif ext in (".md", ".markdown", ".txt"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedDocumentError(
            f"Unsupported file type '{ext or filename}'. Upload a PDF, DOCX, Markdown, or text file."
        )
    return text[:MAX_MARKDOWN_CHARS].strip()


def _pdf_to_markdown(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency declared in requirements
        raise DocumentExtractionError("PDF support requires the 'pypdf' package.") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:  # pypdf raises various read/parse errors
        raise DocumentExtractionError("Could not read this PDF (it may be corrupt or encrypted).") from exc
    return "\n\n".join(p for p in pages if p)


def _docx_to_markdown(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency declared in requirements
        raise DocumentExtractionError("DOCX support requires the 'python-docx' package.") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentExtractionError("Could not read this DOCX (it may be corrupt).") from exc

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style == "title" or style.startswith("heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("heading"):
            lines.append(f"### {text}")
        elif style.startswith("list"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

    return "\n\n".join(lines)
